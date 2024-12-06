from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml import OxmlElement

# Создание документа
doc = Document()

# Настройка обычного текста
paragraph = doc.add_paragraph()
run = paragraph.add_run("В микроконтроллерах ATmega, используемых на платформах Arduino, существует три вида памяти:")
run.font.size = Pt(12)

# Добавление списка с отступом
list_items = [
    "Флеш-память: используется для хранения скетчей.",
    "ОЗУ (SRAM — static random access memory, статическая оперативная память с произвольным доступом): используется для хранения и работы переменных.",
    "EEPROM (энергонезависимая память): используется для хранения постоянной информации."
]
for item in list_items:
    paragraph = doc.add_paragraph(f"o {item}")
    paragraph.paragraph_format.left_indent = Pt(36)  # Отступ для списка
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Дополнительный текст
doc.add_paragraph("Флеш-память и EEPROM являются энергонезависимыми видами памяти (данные сохраняются при отключении питания). ОЗУ является энергозависимой памятью.")

# Создание таблицы
table = doc.add_table(rows=4, cols=5)
table.style = 'Table Grid'

headers = ["", "ATmega168", "ATmega328", "ATmega1280", "ATmega2560"]
table.rows[0].cells[0].text = " "
for i, header in enumerate(headers):
    table.rows[0].cells[i].text = header

data = [
    ["Flash\n(1 кБ flash-памяти занят загрузчиком)", "16 Кбайт", "32 Кбайт", "128 Кбайт", "256 Кбайт"],
    ["SRAM", "1 Кбайт", "2 Кбайт", "8 Кбайт", "8 Кбайт"],
    ["EEPROM", "512 байт", "1024 байта", "4 Кбайт", "4 Кбайт"]
]

for i, row_data in enumerate(data):
    for j, value in enumerate(row_data):
        table.rows[i + 1].cells[j].text = value

# Удаление лишних строк между текстом и таблицей
last_paragraph = doc.paragraphs[-1]
if not last_paragraph.text.strip():
    p_element = last_paragraph._element
    p_element.getparent().remove(p_element)
    p_element._p = p_element._element = None

# Добавление текста курсивом
italic_paragraph = doc.add_paragraph()
italic_run = italic_paragraph.add_run(
    "Память EEPROM, по заявлениям производителя, обладает гарантированным жизненным циклом 100 000 операций записи/стирания и 100 лет хранения данных при температуре 25°С. "
    "Эти данные не распространяются на операции чтения данных из EEPROM — чтение данных не лимитировано. Исходя из этого, нужно проектировать свои скетчи максимально щадящими по отношению к EEPROM."
)
italic_run.italic = True

# Сохранение документа
doc.save("file.docx")
