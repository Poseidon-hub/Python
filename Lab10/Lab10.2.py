from docx import Document
from docx.shared import Inches

# Открытие созданного документа
doc = Document("file.docx")

# Добавление изображения
doc.add_paragraph(" ")  # Для отступа
doc.add_picture('image.jpg', width=Inches(4.0))
doc.add_paragraph("Рисунок 1. Пример изображения", style='Caption')

# Сохранение документа
doc.save("file_image.docx")